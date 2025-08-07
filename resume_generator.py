
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from typing import Dict, List, Any
import re

class ResumeGenerator:
    """Generates tailored resume content and documents"""

    def __init__(self, base_url: str = "http://localhost:11434/api/generate", model_name: str = "gpt-oss:20b"):
        self.base_url = base_url
        self.model_name = model_name

    def tailor_resume(self, parsed_resume: Dict[str, Any],
                     job_analysis: Dict[str, Any],
                     model_name: str = "gpt-oss:20b") -> Dict[str, Any]:
        """Main function to tailor resume content"""

        self.model_name = model_name

        # Generate key themes analysis
        key_themes = self._generate_key_themes(job_analysis)

        # Tailor experience section
        tailored_experience = self._tailor_experience_section(
            parsed_resume.get('experience', []), job_analysis
        )

        # Generate new projects
        tailored_projects = self._generate_tailored_projects(
            parsed_resume, job_analysis
        )

        # Identify missing keywords
        missing_keywords = self._identify_missing_keywords(
            parsed_resume, job_analysis
        )

        # Create Word document
        docx_content = self._create_docx_resume(
            parsed_resume, tailored_experience, tailored_projects, key_themes
        )

        return {
            'key_themes': key_themes,
            'experience': tailored_experience,
            'projects': tailored_projects,
            'missing_keywords': missing_keywords,
            'docx_content': docx_content
        }

    def _call_model(self, prompt: str) -> str:
        """Call the local language model via HTTP."""
        try:
            payload = {"model": self.model_name, "prompt": prompt, "stream": False}
            response = requests.post(self.base_url, json=payload, timeout=60)
            response.raise_for_status()
            data = response.json()
            return data.get("response", "").strip()
        except Exception as e:
            return f"Error communicating with model: {e}"

    def _generate_key_themes(self, job_analysis: Dict[str, Any]) -> str:
        """Generate key themes analysis from job description"""

        prompt = f"""
        Analyze this job description and provide a concise summary of:
        1. The main role responsibilities
        2. Key technical skills required
        3. Business domain/industry context
        4. Company culture indicators

        Job Description Analysis:
        - Company: {job_analysis.get('company_info', {}).get('name', 'Unknown')}
        - Industry: {job_analysis.get('industry', 'General')}
        - Technical Skills: {', '.join(job_analysis.get('technical_skills', [])[:10])}
        - Role Requirements: {job_analysis.get('role_requirements', {})}
        - Job Level: {job_analysis.get('job_level', 'Mid-level')}

        Provide a 2-3 sentence summary focusing on what makes this role unique.
        """

        system_prompt = "You are an expert resume analyst who helps job seekers understand job requirements."
        full_prompt = f"{system_prompt}\n{prompt}"

        return self._call_model(full_prompt)

    def _tailor_experience_section(self, experiences: List[Dict[str, Any]], 
                                 job_analysis: Dict[str, Any]) -> str:
        """Tailor experience section bullets to match job requirements"""

        if not experiences:
            return "No experience section found in original resume."

        tailored_sections = []

        for experience in experiences[:3]:  # Limit to top 3 most recent roles
            tailored_bullets = self._tailor_experience_bullets(experience, job_analysis)

            # Format the experience section
            company = experience.get('company', 'Company Name')
            position = experience.get('position', 'Position Title')
            dates = experience.get('dates', 'Dates')

            section = f"""
**{company}** - {position}
*{dates}*

{tailored_bullets}
"""
            tailored_sections.append(section)

        return "\n".join(tailored_sections)

    def _tailor_experience_bullets(self, experience: Dict[str, Any], 
                                 job_analysis: Dict[str, Any]) -> str:
        """Tailor individual experience bullets using AI"""

        original_bullets = experience.get('bullets', [])
        if not original_bullets:
            original_bullets = [f"Worked as {experience.get('position', 'professional')} at {experience.get('company', 'company')}"]

        # Create context about the target job
        job_context = {
            'technical_skills': job_analysis.get('technical_skills', []),
            'responsibilities': job_analysis.get('role_requirements', {}).get('responsibilities', []),
            'industry': job_analysis.get('industry', 'General'),
            'job_level': job_analysis.get('job_level', 'Mid-level')
        }

        prompt = f"""
        Transform these resume bullets to align with this target job while keeping the same achievements and results. 

        ORIGINAL BULLETS:
        {chr(10).join(['• ' + bullet for bullet in original_bullets])}

        TARGET JOB CONTEXT:
        - Industry: {job_context['industry']}
        - Level: {job_context['job_level']}
        - Key Technical Skills: {', '.join(job_context['technical_skills'][:8])}
        - Key Responsibilities: {', '.join(job_context['responsibilities'][:5])}

        INSTRUCTIONS:
        1. Rewrite each bullet to emphasize skills/experiences that match the target job
        2. Use action verbs that align with the target role's language
        3. Keep the same factual achievements but frame them for this specific opportunity
        4. Include relevant technical terms from the job description where truthful
        5. Maintain quantified results and metrics
        6. Each bullet should start with a unique action verb
        7. Format: • [Action verb] [what you did] [result/impact]

        Provide 3-4 strong bullets that show you already have experience doing what this job requires:
        """

        system_prompt = (
            "You are an expert resume writer who specializes in tailoring experience bullets "
            "to match specific job requirements while maintaining truthfulness."
        )
        full_prompt = f"{system_prompt}\n{prompt}"

        result = self._call_model(full_prompt)
        if result.startswith("Error"):
            # Fallback to original bullets if model call fails
            return "\n".join(['• ' + bullet for bullet in original_bullets])
        return result

    def _generate_tailored_projects(self, parsed_resume: Dict[str, Any], 
                                  job_analysis: Dict[str, Any]) -> str:
        """Generate new relevant projects for the target job"""

        # Get existing skills and tools from resume
        existing_skills = parsed_resume.get('skills', [])
        existing_projects = parsed_resume.get('projects', [])

        # Extract tools and technologies mentioned in resume
        resume_text = parsed_resume.get('original_text', '')
        tools_mentioned = self._extract_tools_from_text(resume_text)

        prompt = f"""
        Generate 1-2 realistic academic/independent projects that demonstrate skills for this target job.

        TARGET JOB INFO:
        - Industry: {job_analysis.get('industry', 'General')}
        - Technical Skills Required: {', '.join(job_analysis.get('technical_skills', [])[:10])}
        - Job Level: {job_analysis.get('job_level', 'Mid-level')}
        - Key Responsibilities: {', '.join(job_analysis.get('role_requirements', {}).get('responsibilities', [])[:3])}

        CANDIDATE'S EXISTING SKILLS/TOOLS:
        {', '.join(existing_skills + tools_mentioned)}

        GUIDELINES:
        1. Each project should be relevant to the target role's domain
        2. Use only technologies/tools the candidate already mentions in their resume
        3. Projects should sound realistic and achievable
        4. Include specific, quantifiable outcomes
        5. Format each project as:
           **Project Title** – Project Type
           Month Year – Month Year
           • [Specific accomplishment with metrics]
           • [Technical implementation detail with results]

        Make projects specific to the {job_analysis.get('industry', 'target')} industry and {job_analysis.get('job_level', 'appropriate')} level.

        Generate 1-2 compelling projects:
        """

        system_prompt = (
            "You are an expert resume writer who creates realistic, relevant project descriptions "
            "that align with job requirements."
        )
        full_prompt = f"{system_prompt}\n{prompt}"

        return self._call_model(full_prompt)

    def _extract_tools_from_text(self, text: str) -> List[str]:
        """Extract tools and technologies mentioned in resume text"""

        common_tools = [
            'Python', 'Java', 'JavaScript', 'SQL', 'R', 'C++', 'C#',
            'React', 'Angular', 'Django', 'Flask', 'Spring', 'Node.js',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes',
            'Git', 'Jenkins', 'Jira', 'Confluence',
            'Tableau', 'Power BI', 'Excel', 'Pandas', 'NumPy',
            'TensorFlow', 'PyTorch', 'Scikit-learn',
            'MySQL', 'PostgreSQL', 'MongoDB', 'Redis'
        ]

        found_tools = []
        text_lower = text.lower()

        for tool in common_tools:
            if tool.lower() in text_lower:
                found_tools.append(tool)

        return found_tools

    def _identify_missing_keywords(self, parsed_resume: Dict[str, Any], 
                                 job_analysis: Dict[str, Any]) -> List[str]:
        """Identify keywords from job that are missing from resume"""

        # Get all text from resume
        resume_text = parsed_resume.get('original_text', '').lower()

        # Collect job keywords
        job_keywords = set()

        if 'technical_skills' in job_analysis:
            job_keywords.update([skill.lower() for skill in job_analysis['technical_skills']])

        if 'soft_skills' in job_analysis:
            job_keywords.update([skill.lower() for skill in job_analysis['soft_skills']])

        if 'keywords' in job_analysis:
            job_keywords.update([kw.lower() for kw in job_analysis['keywords']])

        # Find missing keywords
        missing = []
        for keyword in job_keywords:
            if keyword not in resume_text and len(keyword) > 2:
                missing.append(keyword.title())

        # Return top 15 most important missing keywords
        return missing[:15]

    def _create_docx_resume(self, parsed_resume: Dict[str, Any], 
                          tailored_experience: str, tailored_projects: str,
                          key_themes: str) -> bytes:
        """Create a Word document with the tailored resume"""

        # Create new document
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Add header with contact info
        contact_info = parsed_resume.get('contact_info', {})

        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contact_info.get('name', 'Your Name'))
        name_run.font.size = 16
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact details
        contact_details = []
        if contact_info.get('email'):
            contact_details.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_details.append(contact_info['phone'])

        if contact_details:
            contact_para = doc.add_paragraph(' | '.join(contact_details))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Space

        # Add summary if exists
        summary = parsed_resume.get('summary', '')
        if summary:
            doc.add_heading('Professional Summary', level=2)
            doc.add_paragraph(summary)

        # Add tailored experience
        if tailored_experience and "No experience section found" not in tailored_experience:
            doc.add_heading('Professional Experience', level=2)
            # Parse and format the tailored experience
            self._add_formatted_experience(doc, tailored_experience)

        # Add tailored projects
        if tailored_projects and "Error generating projects" not in tailored_projects:
            doc.add_heading('Projects', level=2)
            self._add_formatted_projects(doc, tailored_projects)

        # Add education
        education = parsed_resume.get('education', [])
        if education:
            doc.add_heading('Education', level=2)
            for edu in education:
                doc.add_paragraph(edu.get('degree', 'Degree information'))

        # Add skills
        skills = parsed_resume.get('skills', [])
        if skills:
            doc.add_heading('Skills', level=2)
            skills_para = doc.add_paragraph()
            skills_text = ', '.join(skills[:20])  # Limit to 20 skills
            skills_para.add_run(skills_text)

        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return doc_buffer.getvalue()

    def _add_formatted_experience(self, doc: Document, experience_text: str):
        """Add formatted experience section to document"""

        # Split by sections (company entries)
        sections = experience_text.strip().split('\n\n')

        for section in sections:
            if section.strip():
                lines = section.strip().split('\n')

                for i, line in enumerate(lines):
                    line = line.strip()
                    if not line:
                        continue

                    if line.startswith('**') and line.endswith('**'):
                        # Company and position line
                        para = doc.add_paragraph()
                        # Remove markdown formatting
                        clean_line = line.replace('**', '')
                        para.add_run(clean_line).bold = True
                    elif line.startswith('*') and line.endswith('*'):
                        # Dates line
                        para = doc.add_paragraph()
                        clean_line = line.replace('*', '')
                        para.add_run(clean_line).italic = True
                    elif line.startswith('•') or line.startswith('-'):
                        # Bullet point
                        para = doc.add_paragraph(line, style='List Bullet')
                    else:
                        # Regular text
                        doc.add_paragraph(line)

                doc.add_paragraph()  # Space between entries

    def _add_formatted_projects(self, doc: Document, projects_text: str):
        """Add formatted projects section to document"""

        # Split by project entries
        lines = projects_text.strip().split('\n')
        current_project = []

        for line in lines:
            line = line.strip()
            if not line:
                if current_project:
                    self._add_single_project(doc, current_project)
                    current_project = []
                continue

            current_project.append(line)

        # Add final project if exists
        if current_project:
            self._add_single_project(doc, current_project)

    def _add_single_project(self, doc: Document, project_lines: List[str]):
        """Add a single project to the document"""

        for i, line in enumerate(project_lines):
            if line.startswith('**') and '–' in line:
                # Project title and type
                para = doc.add_paragraph()
                clean_line = line.replace('**', '')
                para.add_run(clean_line).bold = True
            elif any(month in line for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
                                               'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']):
                # Date line
                para = doc.add_paragraph()
                para.add_run(line).italic = True
            elif line.startswith('•') or line.startswith('-'):
                # Bullet point
                para = doc.add_paragraph(line, style='List Bullet')
            else:
                # Regular description
                doc.add_paragraph(line)

        doc.add_paragraph()  # Space after project
